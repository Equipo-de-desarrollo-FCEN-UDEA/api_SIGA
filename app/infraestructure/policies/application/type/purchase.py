from __future__ import annotations

import os
from tempfile import NamedTemporaryFile

from docx import Document
from fastapi import HTTPException

from app.infraestructure.policies.application.application_flow import ApplicationFlow
from app.schemas.application.type.purchase import PurchaseComplete
from app.services.application.type.purchase import purchase_svc
from app.services.application.user_application_status import user_application_status_svc


class PurchaseFlow(ApplicationFlow):
    def __init__(self, user_application):
        super().__init__(user_application)

    async def complete_information(self, **kwargs):
        db_postgres = kwargs.get('db_postgres')
        purchase = await purchase_svc.get(
            id=self.user_application.id,
            db=kwargs.get('db_mongo'),
        )

        data = kwargs.get('purchase_complete')
        obj_in: PurchaseComplete = PurchaseComplete(**data)

        await purchase_svc.update(
            db_obj=purchase,
            obj_in=obj_in,
            db=kwargs.get('db_mongo'),
        )

        user_application_status = self.get_next_status(
            updated_by=kwargs.get('current_user').id,
            observation=kwargs.get('observation'),
        )

        user_application_status_svc.create(
            obj_in=user_application_status, db=db_postgres,
        )


def generate_format(purchase_dict: dict, path: str):

    purchase_data = {
        'date': f"{purchase_dict['status'][0].date.strftime('%Y-%m-%d')}",
        'academic_unit': str(purchase_dict.get('academic_unit', '')),
        'cost_center': '',
        'need': str(purchase_dict.get('need', '')),
        'description': str(purchase_dict.get('description', '')),
        'responsible_condition': str(purchase_dict.get('responsible_condition', '')),
        'estimated_budget': str(purchase_dict.get('estimated_budget', '')),
        'marco_yes': '☒' if purchase_dict.get('marco_agreement') else '☐',
        'marco_no': '☒' if purchase_dict.get('marco_agreement') is False else '☐',
        'annual_plan': (
            '☒' if (
                hasattr(
                    purchase_dict.get(
                        'prior_consultation', None,
                    ), 'annual_plan',
                )
                and purchase_dict[
                    'prior_consultation'
                ].annual_plan.is_true
            ) else '☐'
        ),
        'code_annual_plan': (
            str(purchase_dict['prior_consultation'].annual_plan.code)
            if hasattr(
                purchase_dict.get(
                    'prior_consultation', None,
                ), 'annual_plan',
            )
            else ''
        ),
        'bank_consultation': (
            '☒' if (
                hasattr(
                    purchase_dict.get(
                        'prior_consultation', None,
                    ), 'bank_consultation',
                )
                and purchase_dict[
                    'prior_consultation'
                ].bank_consultation.is_true
            ) else '☐'
        ),
        'code_bank_consultation': (
            str(
                purchase_dict[
                    'prior_consultation'
                ].bank_consultation.code,
            )
            if hasattr(
                purchase_dict.get(
                    'prior_consultation', None,
                ), 'bank_consultation',
            )
            else ''
        ),
        'contract': (
            str(
                purchase_dict[
                    'prior_consultation'
                ].contract,
            )
            if hasattr(
                purchase_dict.get(
                    'prior_consultation', None,
                ), 'contract',
            )
            else ''
        ),
        'signature': '',
        'signature_delegate': '',
    }

    try:
        document = Document(path)
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f'Error al cargar la plantilla: {str(e)}',
        )

    for paragraph in document.paragraphs:
        for key, value in purchase_data.items():
            if f'{{{{{key}}}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    f'{{{{{key}}}}}', value,
                )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in purchase_data.items():
                    if f'{{{{{key}}}}}' in cell.text:
                        cell.text = cell.text.replace(f'{{{{{key}}}}}', value)

    # Crear archivo temporal para guardar el documento generado
    tmp_file = NamedTemporaryFile(delete=False, suffix='.docx')
    document.save(tmp_file.name)
    tmp_file.close()
    return tmp_file.name


def delete_file(file_path: str):
    """Función para eliminar el archivo después de la descarga."""
    try:
        os.remove(file_path)
    except FileNotFoundError:
        print(f'Archivo no encontrado: {file_path}')

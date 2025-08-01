from __future__ import annotations

import os
from tempfile import NamedTemporaryFile
from uuid import UUID

from docx import Document
from fastapi import HTTPException

from app.core.constants import TYPE_INTERNATIONAL_RELATIONS
from app.infraestructure.policies.application.application_flow import ApplicationFlow
from app.services.organization.academic_unit import academic_unit_svc


class MobilityFlow(ApplicationFlow):
    def __init__(self, user_application):
        super().__init__(user_application)

    async def send_to_international_relations(self, **kwargs):
        db_postgres = kwargs.get('db_postgres')

        # -------------------------------------------------------------
        # Cálculo de la facultad a la que se envía la movilidad:
        # En el endpoint de create_mobility primero se crea en el comité al que va
        # dirigido, y luego a la facultad. Por eso se toma el elemento 1 de la lista
        # de unidades académicas en donde existe la solicitud.

        faculty_id = (
            self.user_application.user_application_academic_units[1]
            .academic_unit_id
        )
        academic_unit = academic_unit_svc.get(
            id=faculty_id,
            db=db_postgres,
        )

        dependencies = academic_unit.academic_units

        for daugther in dependencies:
            if daugther.academic_unit_type.id == UUID(TYPE_INTERNATIONAL_RELATIONS):
                international_relations_id = daugther.id
                break

        response = await self.send_to_academic_unit(
            academic_unit_id=international_relations_id,
            **kwargs,
        )

        return response


def generate_format(mobility_dict: dict, path: str):

    mobility_data = {
        'date': f"{mobility_dict['status'][0].date.strftime('%Y-%m-%d')}",
        'name': f"{mobility_dict['name']} {mobility_dict['last_name']}",
        'phone': mobility_dict['phone'],
        'email': mobility_dict['email'],
        'identification_type': f"""
{mobility_dict['identification_type'].value.replace('_', ' ')}""",
        'identification_number': mobility_dict['identification_number'],
        'nationality': 'HAY QUE PEDIR LA NACIONALIDAD',
        'rol': f"{mobility_dict['student_rol']['name']}",
        'school': mobility_dict['school'],
        'current_academic_program': f"{mobility_dict['current_program']['name']}",
        'semester': 'HAY QUE PEDIR EL SEMESTRE',
        'name_coordinator': 'HAY QUE INVENTARLO',
        'phone_coordinator': 'HAY QUE INVENTARLO',
        'email_coordinator': 'HAY QUE INVENTARLO',
        'incoming_leaving': f"{mobility_dict['type'].split()[0]}",
        'national_international': f"{mobility_dict['type'].split()[1]}",
        'process': f"{mobility_dict['process']}",
        'destination_country': mobility_dict['destination_country'],
        'destination_institution': mobility_dict['destination_institution'],
        'academic_program': mobility_dict['academic_program'],
        'name_contact_person': mobility_dict['name_contact_person'],
        'cellphone_contact_person': mobility_dict['cellphone_contact_person'],
        'email_contact_person': mobility_dict['email_contact_person'],
        'date_start': mobility_dict['date_start'].strftime('%Y-%m-%d'),
        'date_end': mobility_dict['date_end'].strftime('%Y-%m-%d'),
        'code1': mobility_dict.get('code1', ''),
        'subject1': mobility_dict.get('subject1', ''),
        'recognized_code1': mobility_dict.get('recognized_code1', ''),
        'recognized_subject1': mobility_dict.get('recognized_subject1', ''),
        'code2': mobility_dict.get('code2', ''),
        'subject2': mobility_dict.get('subject2', ''),
        'recognized_code2': mobility_dict.get('recognized_code2', ''),
        'recognized_subject2': mobility_dict.get('recognized_subject2', ''),
        'code3': mobility_dict.get('code3', ''),
        'subject3': mobility_dict.get('subject3', ''),
        'recognized_code3': mobility_dict.get('recognized_code3', ''),
        'recognized_subject3': mobility_dict.get('recognized_subject3', ''),
        'code4': mobility_dict.get('code4', ''),
        'subject4': mobility_dict.get('subject4', ''),
        'recognized_code4': mobility_dict.get('recognized_code4', ''),
        'recognized_subject4': mobility_dict.get('recognized_subject4', ''),
        'signature': '',
        'signature_responsible': '',
        'date_report': mobility_dict['date_report'].strftime('%Y-%m-%d'),
    }

    try:
        document = Document(path)
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f'Error al cargar la plantilla: {str(e)}',
        )

    # Reemplazar placeholders en el documento
    for paragraph in document.paragraphs:
        for key, value in mobility_data.items():
            if f'{{{{{key}}}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    f'{{{{{key}}}}}', value,
                )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in mobility_data.items():
                    if f'{{{{{key}}}}}' in cell.text:
                        cell.text = cell.text.replace(f'{{{{{key}}}}}', value)

    # Crear archivo temporal para guardar el documento generado
    tmp_file = NamedTemporaryFile(delete=False, suffix='.docx')
    document.save(tmp_file.name)
    tmp_file.close()
    return tmp_file.name


def delete_file(file_path: str):
    """Función para eliminar el archivo después de la descarga."""
    try:
        os.remove(file_path)
    except FileNotFoundError:
        print(f'Archivo no encontrado: {file_path}')
